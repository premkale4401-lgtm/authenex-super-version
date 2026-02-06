import hashlib
from docx import Document
import PyPDF2

def file_hash(path: str) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        h.update(f.read())
    return h.hexdigest()

def extract_text_from_pdf(path: str) -> str:
    text = ""
    with open(path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            text += page.extract_text() or ""
    return text

def extract_text_from_docx(path: str) -> str:
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs)

def extract_metadata_pdf(path: str) -> dict:
    with open(path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        return reader.metadata or {}

def extract_metadata_docx(path: str) -> dict:
    doc = Document(path)
    return doc.core_properties.__dict__
